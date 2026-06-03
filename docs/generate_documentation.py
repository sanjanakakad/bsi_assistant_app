#!/usr/bin/env python3
"""Generate Word (.docx) and Confluence-ready HTML technical documentation."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent
DIAGRAMS = DOCS_DIR / "diagrams"
OUTPUT_DOCX = DOCS_DIR / "BSI_Assistant_Technical_Documentation.docx"
OUTPUT_HTML = DOCS_DIR / "BSI_Assistant_Technical_Documentation.html"


def set_cell_shading(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def svg_to_png(svg_path: Path) -> Path | None:
    png_path = svg_path.with_suffix(".png")
    ql_png = svg_path.with_name(svg_path.name + ".png")
    if png_path.exists():
        return png_path
    if ql_png.exists():
        ql_png.rename(png_path)
        return png_path

    import shutil
    import subprocess

    if shutil.which("qlmanage"):
        subprocess.run(
            ["qlmanage", "-t", "-s", "1400", "-o", str(svg_path.parent), str(svg_path)],
            check=False,
            capture_output=True,
        )
        if ql_png.exists():
            ql_png.rename(png_path)
            return png_path

    if shutil.which("rsvg-convert"):
        subprocess.run(
            ["rsvg-convert", "-w", "1400", str(svg_path), "-o", str(png_path)],
            check=True,
            capture_output=True,
        )
        return png_path if png_path.exists() else None

    return None


def add_diagram(doc: Document, svg_path: Path, caption: str, width: float = 6.2) -> None:
    png_path = svg_to_png(svg_path)

    if png_path and png_path.exists():
        doc.add_picture(str(png_path), width=Inches(width))
    else:
        note = doc.add_paragraph()
        note.add_run(
            f"[See diagram file: docs/diagrams/{svg_path.name}] "
            f"Open the HTML version or import SVGs into Confluence."
        )
        note.runs[0].italic = True

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)


def build_docx() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("BSI IT-Grundschutz Assistant", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Backend Architecture & Technology Stack")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph("Document version: 1.0  |  Application: bsi_assistant_app_2.py")
    doc.add_paragraph()

    # 1. Executive Summary
    add_heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph(
        "The BSI IT-Grundschutz Assistant is a Streamlit web application that supports "
        "industrial security architects in two complementary workflows: (1) conversational "
        "querying of the BSI Compendium via an OpenAI Assistant with file search, and "
        "(2) deterministic generation of VDI 2182 threat impact matrices from predefined "
        "system architectures."
    )
    doc.add_paragraph(
        "The backend follows a hybrid design: LLM-powered retrieval for flexible Q&A, "
        "combined with rule-based aggregation and structured YAML/JSON knowledge stores "
        "for auditable, repeatable compliance outputs."
    )

    # 2. Architecture
    add_heading(doc, "2. System Architecture", 1)
    add_diagram(doc, DIAGRAMS / "architecture.svg", "Figure 1 — High-level system architecture")

    # 3. Tech Stack
    add_heading(doc, "3. Technology Stack", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Layer", "Technology", "Purpose"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8F0FE")

    rows = [
        ("Frontend / UI", "Streamlit", "Chat interface, sidebar navigation, editable threat matrix (data_editor), CSV export"),
        ("LLM & RAG", "OpenAI Assistants API + File Search", "Grounded Q&A over BSI Compendium PDF hosted in OpenAI vector store"),
        ("Component mapping", "RapidFuzz + Regex", "Map free-text names (PLC, SPS, sensors) to standardized IND modules"),
        ("Data processing", "pandas", "Threat matrix as DataFrame; change detection for expert edits"),
        ("Knowledge store", "YAML + JSON", "Synonyms, architectures, threats, VDI mappings, expert templates"),
        ("Configuration", "python-dotenv / Streamlit Secrets", "API keys and Assistant ID for local and cloud deployment"),
        ("Runtime", "Python 3.10+", "Core application logic"),
    ]
    for layer, tech, purpose in rows:
        row = table.add_row().cells
        row[0].text = layer
        row[1].text = tech
        row[2].text = purpose

    doc.add_paragraph()

    # 4. Data Layer
    add_heading(doc, "4. Data Layer", 1)
    add_diagram(doc, DIAGRAMS / "data_layer.svg", "Figure 2 — Static knowledge files loaded at startup")

    files_table = doc.add_table(rows=1, cols=2)
    files_table.style = "Table Grid"
    fh = files_table.rows[0].cells
    fh[0].text = "File"
    fh[1].text = "Description"
    set_cell_shading(fh[0], "E6F4EA")
    set_cell_shading(fh[1], "E6F4EA")

    for fname, desc in [
        ("synonyms_2022.yaml", "Canonical component names and aliases (e.g. SPS → PLC)"),
        ("ind_modules_2022.yaml", "Authoritative BSI IND module codes and titles"),
        ("component_compositions.yaml", "Predefined architectures mapped to BSI modules"),
        ("bsi_threat_store.json", "Pre-extracted threat titles per IND module with citations"),
        ("vdi2182_threat_mapping.yaml", "Deterministic C / I / A / S impact flags per threat"),
        ("expert_threat_templates.json", "Persisted expert-added or customized threats"),
    ]:
        r = files_table.add_row().cells
        r[0].text = fname
        r[1].text = desc

    # 5. Chat mode
    add_heading(doc, "5. Backend Flow — Chat Mode (Phases 1–3)", 1)
    add_diagram(doc, DIAGRAMS / "chat_sequence.svg", "Figure 3 — Sequence diagram: chat and component mapping flow")

    for step in [
        "User submits a natural-language question or component list via st.chat_input.",
        "split_components() parses input on delimiters (comma, semicolon, newline, 'and').",
        "map_to_ind() resolves components using regex, exact synonym lookup, and fuzzy matching (≥86% score).",
        "If components are found: structured prompts are sent per component to the OpenAI Assistant.",
        "Assistant performs file search on the BSI PDF and returns threat titles with file citations.",
        "If no components are detected: the full question is sent directly to the Assistant.",
        "Responses are assembled and displayed in the chat UI; session state preserves conversation history.",
    ]:
        add_bullet(doc, step)

    # 6. Matrix mode
    add_heading(doc, "6. Backend Flow — Matrix Mode (Phases 4–5)", 1)
    add_diagram(doc, DIAGRAMS / "matrix_flow.svg", "Figure 4 — Deterministic threat matrix pipeline")

    for step in [
        "User selects an architecture from the sidebar (e.g. plc_basic, plc_with_field_devices).",
        "aggregate_threats() collects threats from bsi_threat_store.json for all linked BSI modules.",
        "Duplicate threats are merged; source modules are tracked per row.",
        "VDI impact flags (C, I, A, S) are applied from vdi2182_threat_mapping.yaml — no LLM involved.",
        "Expert templates from expert_threat_templates.json are merged as additional rows.",
        "Matrix is rendered in st.data_editor; edits are diffed against the original for audit labels.",
        "Output: CSV download and optional JSON persistence of expert templates.",
    ]:
        add_bullet(doc, step)

    # 7. Design principles
    add_heading(doc, "7. Key Design Principles", 1)
    principles = [
        ("Hybrid intelligence: ", "LLM for reading/explaining; rules for scoring and aggregation."),
        ("Auditability: ", "Every matrix row tagged as BSI, Modified, or Expert origin."),
        ("Anti-hallucination: ", "VDI flags and architecture mappings are human-reviewed YAML, not LLM-generated."),
        ("Separation of modes: ", "Chat = exploratory; Matrix = structured deliverable."),
    ]
    for prefix, text in principles:
        add_bullet(doc, text, prefix)

    # 8. Deployment
    add_heading(doc, "8. Deployment Notes", 1)
    doc.add_paragraph(
        "The application deploys to Streamlit Cloud via GitHub. Secrets (OPENAI_API_KEY, "
        "OPENAI_ASSISTANT_ID) are configured in Streamlit Secrets — not committed to the repository. "
        "The BSI PDF resides in the OpenAI vector store attached to the Assistant; it is not bundled in the repo."
    )

    add_heading(doc, "9. Limitations", 2)
    for lim in [
        "Chat quality depends on Assistant configuration and PDF coverage.",
        "Expert JSON saves may not persist on Streamlit Cloud (ephemeral filesystem); CSV export is recommended.",
        "Threat store and VDI mappings require periodic human review for production use.",
    ]:
        add_bullet(doc, lim)

    doc.save(OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")


def build_html() -> None:
    diagrams = ["architecture", "data_layer", "chat_sequence", "matrix_flow"]
    svg_blocks = ""
    for name in diagrams:
        svg_content = (DIAGRAMS / f"{name}.svg").read_text(encoding="utf-8")
        svg_blocks += f"""
        <div class="diagram-block">
          <div class="diagram">{svg_content}</div>
          <p class="caption">Figure — {name.replace('_', ' ').title()}</p>
        </div>
        """

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>BSI Assistant — Technical Documentation</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif; max-width: 960px; margin: 40px auto; padding: 0 24px; color: #202124; line-height: 1.6; }}
    h1 {{ color: #1f4e79; border-bottom: 3px solid #1f4e79; padding-bottom: 8px; }}
    h2 {{ color: #1f4e79; margin-top: 32px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #dadce0; padding: 10px 12px; text-align: left; }}
    th {{ background: #e8f0fe; }}
    .diagram-block {{ margin: 24px 0; text-align: center; background: #f8f9fa; padding: 16px; border-radius: 8px; border: 1px solid #dadce0; }}
    .diagram svg {{ max-width: 100%; height: auto; }}
    .caption {{ font-style: italic; color: #5f6368; font-size: 13px; }}
    .info-box {{ background: #e8f0fe; border-left: 4px solid #1f4e79; padding: 12px 16px; margin: 16px 0; }}
    ul {{ padding-left: 24px; }}
    code {{ background: #f1f3f4; padding: 2px 6px; border-radius: 4px; }}
  </style>
</head>
<body>

<h1>BSI IT-Grundschutz Assistant — Backend &amp; Tech Stack</h1>
<p><strong>Application:</strong> <code>bsi_assistant_app_2.py</code> &nbsp;|&nbsp; <strong>Version:</strong> 1.0</p>

<div class="info-box">
  <strong>Summary:</strong> Hybrid Streamlit app combining OpenAI Assistant file search (RAG) for BSI Q&amp;A
  with deterministic YAML/JSON-driven threat matrix generation aligned to VDI 2182 (CIA + Safety).
</div>

<h2>1. System Architecture</h2>
{svg_blocks.split('</div>')[0] + '</div>' if False else ''}
"""
    # Fix - I need to inject svg_blocks properly in the f-string. Let me rewrite build_html more cleanly.
    pass


def build_html_clean() -> None:
    diagram_html = ""
    captions = {
        "architecture": "Figure 1 — High-level system architecture",
        "data_layer": "Figure 2 — Static knowledge files",
        "chat_sequence": "Figure 3 — Chat mode sequence diagram",
        "matrix_flow": "Figure 4 — Matrix mode pipeline",
    }
    for name, caption in captions.items():
        svg = (DIAGRAMS / f"{name}.svg").read_text(encoding="utf-8")
        diagram_html += f'<div class="diagram-block"><div class="diagram">{svg}</div><p class="caption">{caption}</p></div>\n'

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>BSI Assistant — Technical Documentation</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif; max-width: 960px; margin: 40px auto; padding: 0 24px; color: #202124; line-height: 1.6; }}
    h1 {{ color: #1f4e79; border-bottom: 3px solid #1f4e79; padding-bottom: 8px; }}
    h2 {{ color: #1f4e79; margin-top: 32px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #dadce0; padding: 10px 12px; text-align: left; vertical-align: top; }}
    th {{ background: #e8f0fe; }}
    .diagram-block {{ margin: 24px 0; text-align: center; background: #f8f9fa; padding: 16px; border-radius: 8px; border: 1px solid #dadce0; overflow-x: auto; }}
    .diagram svg {{ max-width: 100%; height: auto; }}
    .caption {{ font-style: italic; color: #5f6368; font-size: 13px; }}
    .info-box {{ background: #e8f0fe; border-left: 4px solid #1f4e79; padding: 12px 16px; margin: 16px 0; }}
    ul {{ padding-left: 24px; }}
    code {{ background: #f1f3f4; padding: 2px 6px; border-radius: 4px; font-size: 90%; }}
    .confluence-tip {{ background: #fef7e0; border: 1px solid #f9ab00; padding: 12px 16px; border-radius: 6px; margin-top: 32px; }}
  </style>
</head>
<body>

<h1>BSI IT-Grundschutz Assistant — Backend &amp; Tech Stack</h1>
<p><strong>Application:</strong> <code>bsi_assistant_app_2.py</code> &nbsp;|&nbsp; <strong>Version:</strong> 1.0</p>

<div class="info-box">
  <strong>Summary:</strong> Hybrid Streamlit app combining OpenAI Assistant file search (RAG) for BSI Q&amp;A
  with deterministic YAML/JSON-driven threat matrix generation aligned to VDI 2182 (CIA + Safety).
</div>

<h2>1. System Architecture</h2>
{diagram_html.split('<div class="diagram-block">')[1].split('</div>')[0] if False else ''}
"""
    # Still messy - let me just write the full HTML file directly with write tool instead of this broken f-string approach
    OUTPUT_HTML.write_text(get_html_content(diagram_html), encoding="utf-8")
    print(f"Created: {OUTPUT_HTML}")


def get_html_content(diagram_html: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8"/>
  <title>BSI Assistant — Technical Documentation</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif; max-width: 960px; margin: 40px auto; padding: 0 24px; color: #202124; line-height: 1.6; }}
    h1 {{ color: #1f4e79; border-bottom: 3px solid #1f4e79; padding-bottom: 8px; }}
    h2 {{ color: #1f4e79; margin-top: 32px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #dadce0; padding: 10px 12px; text-align: left; vertical-align: top; }}
    th {{ background: #e8f0fe; }}
    .diagram-block {{ margin: 24px 0; text-align: center; background: #f8f9fa; padding: 16px; border-radius: 8px; border: 1px solid #dadce0; overflow-x: auto; }}
    .diagram svg {{ max-width: 100%; height: auto; }}
    .caption {{ font-style: italic; color: #5f6368; font-size: 13px; }}
    .info-box {{ background: #e8f0fe; border-left: 4px solid #1f4e79; padding: 12px 16px; margin: 16px 0; }}
    ul {{ padding-left: 24px; }}
    code {{ background: #f1f3f4; padding: 2px 6px; border-radius: 4px; }}
    .confluence-tip {{ background: #fef7e0; border: 1px solid #f9ab00; padding: 12px 16px; border-radius: 6px; margin-top: 32px; }}
  </style>
</head>
<body>

<h1>BSI IT-Grundschutz Assistant — Backend &amp; Tech Stack</h1>
<p><strong>Application:</strong> <code>bsi_assistant_app_2.py</code> &nbsp;|&nbsp; <strong>Version:</strong> 1.0</p>

<div class="info-box">
  <strong>Summary:</strong> Hybrid Streamlit app combining OpenAI Assistant file search (RAG) for BSI Q&amp;A
  with deterministic YAML/JSON-driven threat matrix generation aligned to VDI 2182 (CIA + Safety).
</div>

<h2>1. System Architecture</h2>
{diagram_html}

<h2>2. Technology Stack</h2>
<table>
  <tr><th>Layer</th><th>Technology</th><th>Purpose</th></tr>
  <tr><td>Frontend / UI</td><td>Streamlit</td><td>Chat, sidebar, editable matrix, CSV export</td></tr>
  <tr><td>LLM &amp; RAG</td><td>OpenAI Assistants API + File Search</td><td>Grounded Q&amp;A over BSI PDF in vector store</td></tr>
  <tr><td>Component mapping</td><td>RapidFuzz + Regex</td><td>Free-text → standardized IND modules</td></tr>
  <tr><td>Data processing</td><td>pandas</td><td>Threat matrix DataFrame; edit diff detection</td></tr>
  <tr><td>Knowledge store</td><td>YAML + JSON</td><td>Synonyms, architectures, threats, VDI mappings</td></tr>
  <tr><td>Configuration</td><td>python-dotenv / Streamlit Secrets</td><td>API keys (local + cloud)</td></tr>
  <tr><td>Runtime</td><td>Python 3.10+</td><td>Application logic</td></tr>
</table>

<h2>3. Chat Mode — Backend Steps (Phases 1–3)</h2>
<ol>
  <li>User input via <code>st.chat_input</code></li>
  <li><code>split_components()</code> — parse delimiters</li>
  <li><code>map_to_ind()</code> — regex + synonym lookup + fuzzy match (≥86%)</li>
  <li>OpenAI thread created; Assistant run with file search</li>
  <li>Response with threat titles and PDF citations rendered in chat</li>
</ol>

<h2>4. Matrix Mode — Backend Steps (Phases 4–5)</h2>
<ol>
  <li>Architecture selected from <code>component_compositions.yaml</code></li>
  <li><code>aggregate_threats()</code> loads threats from JSON per BSI module</li>
  <li>VDI C/I/A/S flags applied from YAML mapping (deterministic)</li>
  <li>Expert templates merged; matrix shown in <code>st.data_editor</code></li>
  <li>Audit labels: 🟢 BSI | 🟡 Modified | 🔵 Expert → CSV export</li>
</ol>

<h2>5. Design Principles</h2>
<ul>
  <li><strong>Hybrid AI:</strong> LLM for reading; rules for compliance outputs</li>
  <li><strong>Auditability:</strong> Origin and modification status on every row</li>
  <li><strong>Anti-hallucination:</strong> VDI flags are human-reviewed YAML, not LLM-generated</li>
</ul>

<div class="confluence-tip">
  <strong>Confluence import tip:</strong> Copy this entire page into Confluence using <em>Paste from Word/HTML</em>,
  or attach the <code>.docx</code> file. SVG diagrams can also be imported individually from
  <code>docs/diagrams/</code> using Confluence's diagram tool or as image attachments.
</div>

</body>
</html>"""


if __name__ == "__main__":
    diagram_html = ""
    captions = {
        "architecture": "Figure 1 — High-level system architecture",
        "data_layer": "Figure 2 — Static knowledge files",
        "chat_sequence": "Figure 3 — Chat mode sequence diagram",
        "matrix_flow": "Figure 4 — Matrix mode pipeline",
    }
    for name, caption in captions.items():
        svg = (DIAGRAMS / f"{name}.svg").read_text(encoding="utf-8")
        diagram_html += (
            f'<div class="diagram-block"><div class="diagram">{svg}</div>'
            f'<p class="caption">{caption}</p></div>\n'
        )

    build_docx()
    OUTPUT_HTML.write_text(get_html_content(diagram_html), encoding="utf-8")
    print(f"Created: {OUTPUT_HTML}")
